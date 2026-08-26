# -*- coding: utf-8 -*-
"""Genera el .docx del Taller Integrador completado con el proyecto Sistema de Reservas."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\borgo\OneDrive\Escritorio\Proyecto T. Integrador\taller-reajuste-COMPLETADO.docx"

INDIGO     = RGBColor(0x4F, 0x46, 0xE5)
INDIGO_DK  = RGBColor(0x37, 0x30, 0xA3)
INDIGO_BG  = "EEF2FF"
SLATE_BG   = "F8FAFC"
TEXT_DARK  = RGBColor(0x1A, 0x1D, 0x2E)

doc = Document()

# márgenes
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

# fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(4)

def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = INDIGO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_p(text, bold_label=None, justify=True):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label); r.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# ── ENCABEZADO ──
add_h1("Taller Integrador")
add_h2("Actividad de Ajuste y Revisión Integral del Proyecto")

add_p("Ignacio Borgonovo", bold_label="Nombre y Apellido:", justify=False)
add_p("Sistema de Reservas - Gestión de Materiales", bold_label="Proyecto:", justify=False)
add_p("17/06/2026", bold_label="Fecha:", justify=False)

doc.add_paragraph()

# ── OBJETIVO ──
add_h3("Objetivo")
add_p("Realizar una revisión integral del proyecto desarrollado hasta el momento, "
      "detectando errores, inconsistencias y oportunidades de mejora para dejar "
      "la documentación preparada para la evaluación parcial.")

# ── CONSIGNA 1 ──
add_h2("1. Revisión de la documentación")
add_p("Se analizó cada apartado del proyecto. La tabla resume el estado, los problemas "
      "detectados y las correcciones aplicadas en esta iteración.")

filas = [
    ("Identificación del problema", "Sí",
     "La descripción inicial era muy genérica: 'gestionar materiales'.",
     "Se reformuló centrándose en el problema concreto de la escuela: evitar superposición "
     "de reservas de equipos informáticos y talleres, y dar visibilidad al stock en tiempo real."),
    ("Relevamiento", "Sí",
     "Faltaba detalle sobre la cantidad real de equipos y los talleres disponibles.",
     "Se relevó el inventario: 20 notebooks, 10 routers WiFi, 8 proyectores, Taller A (30 personas) "
     "y Taller B (25 personas). Se descartaron las PCs de escritorio porque no existen en la escuela."),
    ("Actores", "Sí",
     "Solo se contemplaba 'el docente'. Faltaba el rol de gestión.",
     "Se definieron dos actores: Docente (reserva y consulta) y Coordinador (gestiona inventario "
     "y aprueba). Se sumó login con usuario y contraseña por rol."),
    ("Requerimientos", "Sí",
     "Requerimientos no funcionales sin documentar.",
     "Se agregaron requerimientos no funcionales: interfaz responsive (escritorio y celular), "
     "autenticación, estadísticas de uso por día y semana e indicadores visuales de disponibilidad."),
    ("Diseño funcional", "Sí",
     "La interfaz original era plana y poco clara.",
     "Se rediseñó la UI con tres pestañas (Panel, Reservas, Stock), tarjetas KPI, anillos de "
     "progreso para items disponibles y un gráfico de barras semanal con eje X/Y y valores visibles."),
    ("Flujos", "Sí",
     "El flujo de reserva no estaba definido paso a paso.",
     "Se documentó el flujo: Login → Panel → Reservas → Seleccionar material → Elegir fecha → "
     "Confirmar. Cierre con Logout."),
    ("Modelo de datos - Entidades", "Sí",
     "No estaban claras las entidades principales.",
     "Se definieron: Usuario, Material (con tipo y stock), Reserva (fecha, observaciones), "
     "Categoría y Sesión."),
    ("Modelo de datos - Relaciones", "Sí",
     "Faltaban las cardinalidades.",
     "Usuario 1—N Reserva; Reserva N—1 Material; Material N—1 Categoría. Restricción: no puede "
     "haber dos reservas del mismo material en la misma fecha."),
]

tabla = doc.add_table(rows=1, cols=4)
tabla.style = 'Light Grid Accent 1'
hdr = tabla.rows[0].cells
for i, h in enumerate(["Elemento del proyecto", "¿Completo?", "Problemas detectados", "Correcciones realizadas"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.color.rgb = INDIGO_DK
    set_cell_bg(hdr[i], INDIGO_BG)
    hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for idx, (elem, comp, prob, corr) in enumerate(filas):
    row = tabla.add_row().cells
    for c in row: c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    row[0].text = elem
    row[1].text = comp
    row[2].text = prob
    row[3].text = corr
    if idx % 2 == 1:
        for c in row: set_cell_bg(c, SLATE_BG)

# anchos aproximados
widths = [Cm(4.2), Cm(1.8), Cm(5.0), Cm(6.0)]
for row in tabla.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

doc.add_paragraph()

# ── CONSIGNA 2 ──
add_h2("2. Análisis crítico")

preguntas = [
    ("1. ¿El problema planteado al inicio sigue siendo el mismo o fue necesario redefinirlo?",
     "El problema central se mantuvo —gestionar las reservas de materiales informáticos y "
     "talleres— pero el alcance se redefinió. Originalmente se hablaba en abstracto de "
     "'materiales'; en esta revisión se acotó a los recursos reales de la escuela (notebooks, "
     "routers, proyectores y dos talleres), descartando las PCs de escritorio porque no existen "
     "en la institución."),
    ("2. ¿Los requerimientos cubren todas las necesidades de los usuarios?",
     "Sí. Se cubrió el alta de reservas, la consulta de stock en tiempo real, el control por "
     "rol (Docente / Coordinador) y la visualización de estadísticas de uso por día y por "
     "semana. Como mejora se agregó la diferenciación entre 'libres', 'en uso' y 'capacidad' "
     "(esta última para los talleres)."),
    ("3. ¿Existen funcionalidades que no fueron contempladas inicialmente?",
     "Sí. La revisión incorporó: estadísticas de uso con gráfico semanal interactivo, filtro por "
     "item, indicadores visuales (anillos de progreso, badges de estado) y un diseño responsive "
     "específico para celular con tres breakpoints (820px, 600px, 380px)."),
    ("4. ¿Las entidades y relaciones representan correctamente la información que manejará el sistema?",
     "Sí. Las entidades Usuario, Material, Reserva, Categoría y Sesión cubren el dominio. Las "
     "cardinalidades quedaron explícitas y se sumó la restricción de unicidad por material + "
     "fecha, que evita superposiciones —el problema original que motivó el sistema."),
    ("5. ¿Qué aspecto del proyecto consideran más débil y por qué?",
     "El backend persistente. Por ahora la aplicación funciona con datos cargados en el frontend "
     "(HTML/CSS/JS) y un endpoint PHP de prueba (procesar.php). Falta conectar una base de datos "
     "real (MySQL) para que las reservas y el stock se persistan entre sesiones."),
    ("6. ¿Qué mejoras implementarían antes de comenzar el desarrollo definitivo?",
     "Conectar el frontend a una API PHP + MySQL para almacenar reservas y usuarios, agregar "
     "validación de fechas (no permitir reservas en el pasado ni solapadas), incorporar "
     "notificaciones cuando un item queda con poco stock y separar el panel del Coordinador "
     "con vista de administración de inventario."),
]
for q, a in preguntas:
    p = doc.add_paragraph()
    r = p.add_run(q); r.bold = True
    p.paragraph_format.space_before = Pt(8)
    add_p(a)

# ── CONSIGNA 3 ──
add_h2("3. Ajustes realizados")
add_p("Se detallan los cinco ajustes principales aplicados al proyecto en esta iteración.")

ajustes = [
    ("1", "Rediseño visual profesional del sistema",
     "Mejorar la legibilidad y la presentación institucional del sistema para la evaluación."),
    ("2", "Estadísticas por item con anillos de progreso",
     "Brindar información de disponibilidad de un vistazo y reducir la carga cognitiva del usuario."),
    ("3", "Gráfico de uso semanal con filtros interactivos",
     "Hacer las estadísticas de uso comprensibles y útiles para la toma de decisiones del Coordinador."),
    ("4", "Eliminación de PCs de Escritorio (67 → 52 unidades; 6 → 5 categorías)",
     "Ajustar el sistema al inventario real de la escuela."),
    ("5", "Adaptación responsive para celular (3 breakpoints)",
     "Permitir que docentes y coordinadores usen el sistema desde el celular durante la jornada."),
]
tabla2 = doc.add_table(rows=1, cols=3)
tabla2.style = 'Light Grid Accent 1'
h2 = tabla2.rows[0].cells
for i, h in enumerate(["N°", "Ajuste realizado", "Motivo"]):
    h2[i].text = ""
    r = h2[i].paragraphs[0].add_run(h); r.bold = True; r.font.color.rgb = INDIGO_DK
    set_cell_bg(h2[i], INDIGO_BG)
for idx, (n, a, m) in enumerate(ajustes):
    row = tabla2.add_row().cells
    row[0].text = n
    row[1].text = a
    row[2].text = m
    if idx % 2 == 1:
        for c in row: set_cell_bg(c, SLATE_BG)
widths2 = [Cm(1.2), Cm(7.5), Cm(8.3)]
for row in tabla2.rows:
    for i, c in enumerate(row.cells):
        c.width = widths2[i]

doc.add_paragraph()

# ── CONSIGNA 4 ──
add_h2("4. Reflexión final")
add_p("El proyecto se encuentra en una etapa avanzada de su fase de diseño y prototipo "
      "funcional. Se cuenta con una interfaz completa (login, panel principal, reservas y "
      "stock) implementada en HTML, CSS y JavaScript, y con un endpoint de prueba en PHP. "
      "Durante esta revisión los cambios más importantes fueron el rediseño visual profesional, "
      "la incorporación de estadísticas por día y semana con un gráfico de barras claro y "
      "filtrable, la sustitución de los listados planos por tarjetas con anillos de progreso, "
      "el ajuste del inventario al material real de la escuela y la adaptación completa para "
      "celular mediante media queries en tres breakpoints. La principal dificultad fue iterar "
      "el diseño hasta encontrar una forma de mostrar las estadísticas que sea tanto profesional "
      "como entendible: el sparkline inicial era confuso y hubo que reemplazarlo por un gráfico "
      "con ejes y valores visibles. También fue necesario reorganizar los totales tras eliminar "
      "las PCs de escritorio para mantener coherencia entre todas las vistas. Para la próxima "
      "etapa quedan pendientes: conectar el frontend con una base de datos MySQL a través de "
      "PHP, implementar la validación real de fechas y disponibilidad en el servidor, construir "
      "el panel de administración del Coordinador y agregar notificaciones para los items con "
      "bajo stock. Con estos ajustes el proyecto queda listo para pasar de prototipo a "
      "aplicación operativa.")

# ── ENTREGA ──
add_h3("Entrega y criterios de evaluación")
add_p("documentación revisada y corregida, esta guía completa y las modificaciones señaladas "
      "en la tabla de la Consigna 1.", bold_label="Contenido entregado:")
add_p("completitud de la documentación, coherencia entre problema, requerimientos y diseño, "
      "calidad y fundamentación de los ajustes, y presentación ordenada de la carpeta del proyecto.",
      bold_label="Criterios cubiertos:")

doc.save(OUT)
print("OK ->", OUT)
